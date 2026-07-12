# backend/routes/tools_routes.py
from flask import Blueprint, request, jsonify, send_file
import json
import os
import uuid
import tempfile
import zipfile
from io import BytesIO
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
from config import get_db
from middleware.auth import token_required
from middleware.validation import sanitize_input
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from PIL import Image, ImageDraw, ImageFont
import qrcode
from io import BytesIO
import base64
import re
from docx import Document
import openpyxl
import shutil
import hashlib
import magic
import mimetypes

# Create blueprint WITHOUT url_prefix
tools_bp = Blueprint('tools_bp', __name__)
tools_bp.strict_slashes = False

# ============================================
# ✅ GLOBAL VARIABLES - MUST BE AT TOP
# ============================================

usage_tracking = {}

# ============================================
# 🔒 SECURITY CONFIGURATIONS
# ============================================

# File size limits (in bytes)
MAX_FILE_SIZE = {
    'default': 10 * 1024 * 1024,      # 10MB
    'pdf': 10 * 1024 * 1024,           # 10MB
    'image': 10 * 1024 * 1024,         # 10MB
    'batch': 15 * 1024 * 1024,         # 15MB (for batch operations)
}

# Allowed file extensions and MIME types
ALLOWED_EXTENSIONS = {
    'pdf': '.pdf',
    'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff'],
    'all': ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff']
}

ALLOWED_MIME_TYPES = {
    'pdf': 'application/pdf',
    'image': [
        'image/jpeg', 'image/png', 'image/gif', 'image/bmp', 
        'image/webp', 'image/tiff', 'image/x-ms-bmp'
    ]
}

# ============================================
# 🔒 SECURITY FUNCTIONS
# ============================================

def validate_file_type(file, expected_type='pdf'):
    """
    Validate file type using both extension and MIME type
    Returns: (is_valid, error_message)
    """
    try:
        # Check file extension
        filename = file.filename or ''
        ext = os.path.splitext(filename)[1].lower()
        
        if expected_type == 'pdf':
            if ext != '.pdf':
                return False, "Only PDF files are allowed"
        
        elif expected_type == 'image':
            if ext not in ALLOWED_EXTENSIONS['image']:
                return False, "Only image files are allowed (JPG, PNG, GIF, BMP, WEBP, TIFF)"
        
        # Check MIME type using python-magic
        file_bytes = file.read(1024)
        file.seek(0)  # Reset file pointer
        
        mime = magic.from_buffer(file_bytes, mime=True)
        
        expected_mimes = ALLOWED_MIME_TYPES.get(expected_type, [])
        if isinstance(expected_mimes, str):
            expected_mimes = [expected_mimes]
        
        if expected_type == 'pdf' and mime != 'application/pdf':
            return False, "Invalid PDF file (MIME type mismatch)"
        
        if expected_type == 'image' and mime not in expected_mimes:
            return False, f"Invalid image file (MIME type: {mime})"
        
        return True, ""
        
    except Exception as e:
        print(f"File validation error: {str(e)}")
        # Fallback: check extension only
        return True, ""

def validate_file_size(file, max_size=None):
    """
    Validate file size
    """
    max_size = max_size or MAX_FILE_SIZE['default']
    
    # Get file size
    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    
    if size > max_size:
        return False, f"File size exceeds {max_size // (1024*1024)}MB limit (Current: {size // (1024*1024)}MB)"
    
    return True, ""

def detect_malicious_content(file_bytes):
    """
    Detect potentially malicious content in files
    """
    try:
        # Check for executable signatures
        executable_signatures = [
            b'MZ',  # Windows EXE
            b'\x7fELF',  # Linux ELF
            b'PK\x03\x04',  # ZIP (could be malicious)
            b'%PDF',  # PDF (valid)
        ]
        
        # Check first 50 bytes for suspicious patterns
        header = file_bytes[:50]
        
        # Check for script injection patterns
        suspicious_patterns = [
            b'<script>', b'javascript:', b'data:text/html',
            b'exec(', b'eval(', b'system(', b'passthru(',
            b'<?php', b'<%%', b'<%', b'<jsp:'
        ]
        
        for pattern in suspicious_patterns:
            if pattern in header:
                return True, f"Potentially malicious content detected (pattern: {pattern.decode('utf-8', errors='ignore')})"
        
        # Check for double extensions (e.g., file.pdf.exe)
        return False, ""
        
    except Exception as e:
        print(f"Malicious content detection error: {str(e)}")
        return False, ""

def sanitize_filename(filename):
    """
    Sanitize filename to prevent path traversal attacks
    """
    # Remove path traversal attempts
    filename = filename.replace('../', '').replace('..\\', '')
    filename = filename.replace('/../', '').replace('\\..\\', '')
    
    # Remove dangerous characters
    dangerous_chars = [';', '&', '|', '`', '$', '(', ')', '<', '>', '{', '}', '[', ']', '*', '?']
    for char in dangerous_chars:
        filename = filename.replace(char, '')
    
    # Use secure_filename from werkzeug
    filename = secure_filename(filename)
    
    # If filename becomes empty, generate a random one
    if not filename:
        filename = f"file_{uuid.uuid4().hex[:8]}"
    
    return filename

def cleanup_temp_files(file_paths):
    """
    Safely cleanup temporary files
    """
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception as e:
            print(f"Cleanup error for {path}: {str(e)}")

def get_client_ip(request):
    """
    Get client IP address safely
    """
    # Check for proxy headers
    forwarded = request.headers.get('X-Forwarded-For')
    if forwarded:
        # Get the first IP in the list (client IP)
        ips = forwarded.split(',')
        ip = ips[0].strip()
        # Validate IP format (basic check)
        if re.match(r'^(\d{1,3}\.){3}\d{1,3}$', ip):
            return ip
    
    return request.remote_addr or '0.0.0.0'

def is_rate_limited(ip, tool_name, limit=10, window=60):
    """
    Rate limiting: Max 'limit' requests per 'window' seconds
    """
    key = f"rate_limit:{tool_name}:{ip}"
    current_time = datetime.now().timestamp()
    
    # Clean old entries (if too many)
    if key in usage_tracking and 'requests' in usage_tracking[key]:
        # Remove requests older than window
        cutoff = current_time - window
        usage_tracking[key]['requests'] = [
            t for t in usage_tracking[key]['requests'] 
            if t > cutoff
        ]
        
        # Check if limit exceeded
        if len(usage_tracking[key]['requests']) >= limit:
            return True
    
    return False

def track_request(ip, tool_name):
    """
    Track request for rate limiting
    """
    key = f"rate_limit:{tool_name}:{ip}"
    current_time = datetime.now().timestamp()
    
    if key not in usage_tracking:
        usage_tracking[key] = {'requests': []}
    
    if 'requests' not in usage_tracking[key]:
        usage_tracking[key]['requests'] = []
    
    usage_tracking[key]['requests'].append(current_time)

# ============================================
# WATERMARK CONFIGURATION
# ============================================

WATERMARK_TEXT = "Made with ❤️ by Krynova Technologies"
WATERMARK_FOOTER = "\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\nMade with ❤️ by Krynova Technologies\nVisit: https://krynovatechnology.pythonanywhere.com\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"

def should_add_watermark(is_premium):
    """Determine if watermark should be added"""
    return not is_premium

def add_watermark_to_text(content, is_premium):
    """Add watermark footer to text content"""
    if is_premium:
        return content
    return content + WATERMARK_FOOTER

def add_watermark_to_image(image, is_premium):
    """Add watermark to image using PIL"""
    if is_premium:
        return image
    
    try:
        draw = ImageDraw.Draw(image)
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        except:
            font = ImageFont.load_default()
        
        watermark_text = "Made with ❤️ by Krynova"
        bbox = draw.textbbox((0, 0), watermark_text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        x = (image.width - text_width) // 2
        y = image.height - text_height - 20
        
        # Semi-transparent background
        overlay = Image.new('RGBA', image.size, (255, 255, 255, 0))
        overlay_draw = ImageDraw.Draw(overlay)
        overlay_draw.rectangle(
            [x-10, y-10, x+text_width+10, y+text_height+10],
            fill=(255, 255, 255, 180)
        )
        image = Image.alpha_composite(image.convert('RGBA'), overlay)
        
        draw = ImageDraw.Draw(image)
        draw.text((x, y), watermark_text, fill=(80, 80, 80, 200), font=font)
        
        return image
    except Exception as e:
        print(f"Watermark error: {str(e)}")
        return image

# ============================================
# ✅ USAGE TRACKING FUNCTIONS
# ============================================

def get_usage_count(tool_name, ip_address):
    """Get usage count for a specific tool and IP"""
    key = f"{tool_name}:{ip_address}"
    today = datetime.now().strftime('%Y-%m-%d')
    
    if key not in usage_tracking:
        usage_tracking[key] = {'date': today, 'count': 0}
    
    if usage_tracking[key]['date'] != today:
        usage_tracking[key] = {'date': today, 'count': 0}
    
    return usage_tracking[key]['count']

def increment_usage(tool_name, ip_address):
    """Increment usage count for a specific tool and IP"""
    key = f"{tool_name}:{ip_address}"
    today = datetime.now().strftime('%Y-%m-%d')
    
    if key not in usage_tracking:
        usage_tracking[key] = {'date': today, 'count': 0}
    
    if usage_tracking[key]['date'] != today:
        usage_tracking[key] = {'date': today, 'count': 0}
    
    usage_tracking[key]['count'] += 1
    return usage_tracking[key]['count']

def create_zip_download(files_data, zip_name="download.zip"):
    """Create a ZIP file from multiple files"""
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, data in files_data:
            zip_file.writestr(filename, data)
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# ============================================
# TEST ROUTE
# ============================================

@tools_bp.route('/test', methods=['GET', 'OPTIONS'])
def test_tools():
    """Test route to verify tools blueprint is working"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        return response
    
    return jsonify({
        'success': True,
        'message': 'Tools API is working!',
        'available_tools': [
            '/resume-builder (POST)',
            '/cover-letter (POST)',
            '/qr-generator (POST)',
            '/pdf-to-image (POST)',
            '/pdf-to-word (POST)',
            '/pdf-to-excel (POST)',
            '/image-to-pdf (POST)',
            '/pdf-compressor (POST)',
            '/merge-pdf (POST)',
            '/split-pdf (POST)',
            '/image-resizer (POST)',
            '/text-to-pdf (POST)'
        ]
    })

# ============================================
# 1. RESUME BUILDER
# ============================================

@tools_bp.route('/resume-builder', methods=['POST', 'OPTIONS'])
def build_resume():
    """Generate ATS-friendly resume (Free: 3 per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        ip_address = get_client_ip(request)
        
        # Rate limiting: 20 requests per minute
        if is_rate_limited(ip_address, 'resume_builder', limit=20, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'resume_builder')
        
        usage_count = get_usage_count('resume_builder', ip_address)
        is_premium = data.get('is_premium', False)
        
        if not is_premium and usage_count >= 3:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 3
            }), 403
        
        # Validate required fields
        required = ['name', 'email', 'skills']
        for field in required:
            if not data.get(field):
                return jsonify({'success': False, 'error': f'Missing field: {field}'}), 400
        
        # Sanitize inputs
        name = sanitize_input(data.get('name', ''))
        email = sanitize_input(data.get('email', ''))
        
        # Generate resume content with enhanced function
        resume_content = generate_resume_content(data)
        
        # ✅ Add watermark for free users
        resume_content = add_watermark_to_text(resume_content, is_premium)
        
        increment_usage('resume_builder', ip_address)
        
        return jsonify({
            'success': True,
            'resume': resume_content,
            'usage_count': get_usage_count('resume_builder', ip_address),
            'remaining_free': max(0, 3 - get_usage_count('resume_builder', ip_address)),
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"Resume Builder Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

# ============================================
# 2. COVER LETTER GENERATOR
# ============================================

@tools_bp.route('/cover-letter', methods=['POST', 'OPTIONS'])
def generate_cover_letter():
    """Generate cover letter (Free: 3 per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'cover_letter', limit=20, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'cover_letter')
        
        usage_count = get_usage_count('cover_letter', ip_address)
        is_premium = data.get('is_premium', False)
        
        if not is_premium and usage_count >= 3:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 3
            }), 403
        
        required = ['name', 'position', 'company', 'skills']
        for field in required:
            if not data.get(field):
                return jsonify({'success': False, 'error': f'Missing field: {field}'}), 400
        
        cover_letter = generate_cover_letter_content(data)
        
        # ✅ Add watermark for free users
        cover_letter = add_watermark_to_text(cover_letter, is_premium)
        
        increment_usage('cover_letter', ip_address)
        
        return jsonify({
            'success': True,
            'cover_letter': cover_letter,
            'usage_count': get_usage_count('cover_letter', ip_address),
            'remaining_free': max(0, 3 - get_usage_count('cover_letter', ip_address)),
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"Cover Letter Error: {str(e)}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

# ============================================
# 3. QR CODE GENERATOR
# ============================================

@tools_bp.route('/qr-generator', methods=['POST', 'OPTIONS'])
def generate_qr():
    """Generate QR Code (Free: 5 per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'qr_generator', limit=30, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'qr_generator')
        
        usage_count = get_usage_count('qr_generator', ip_address)
        is_premium = data.get('is_premium', False)
        
        if not is_premium and usage_count >= 5:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 5
            }), 403
        
        content = data.get('content', '')
        if not content:
            return jsonify({'success': False, 'error': 'Content is required'}), 400
        
        # Sanitize content - limit length
        content = content[:5000]  # Max 5000 chars
        
        # Get style and size from request
        style = data.get('style', 'default')
        size = min(int(data.get('size', 250)), 1000)  # Max 1000px
        
        # Generate QR code
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECTION_H,
            box_size=10,
            border=4,
        )
        qr.add_data(content)
        qr.make(fit=True)
        
        # Get colors based on style
        color_map = {
            'default': ('black', 'white'),
            'blue': ('#2563eb', 'white'),
            'green': ('#16a34a', 'white'),
            'purple': ('#7c3aed', 'white'),
            'red': ('#dc2626', 'white'),
            'dark': ('#1f2937', '#f3f4f6'),
        }
        fill_color, back_color = color_map.get(style, ('black', 'white'))
        
        img = qr.make_image(fill_color=fill_color, back_color=back_color)
        
        # Resize if needed
        if size and size != 250:
            img = img.resize((size, size), Image.Resampling.LANCZOS)
        
        # ✅ Add watermark for free users
        if should_add_watermark(is_premium):
            img = add_watermark_to_image(img, is_premium)
        
        # Convert to base64
        buffered = BytesIO()
        img.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
        
        increment_usage('qr_generator', ip_address)
        
        return jsonify({
            'success': True,
            'qr_code': f"data:image/png;base64,{img_base64}",
            'usage_count': get_usage_count('qr_generator', ip_address),
            'remaining_free': max(0, 5 - get_usage_count('qr_generator', ip_address)),
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"QR Generator Error: {str(e)}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

# ============================================
# 4. PDF TO IMAGE - ENHANCED
# ============================================

@tools_bp.route('/pdf-to-image', methods=['POST', 'OPTIONS'])
def pdf_to_image():
    """Convert PDF to Image (Free: 3 pages per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    try:
        file = request.files['file']
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'pdf_to_image', limit=15, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'pdf_to_image')
        
        # ✅ Security: Validate file type
        is_valid, error = validate_file_type(file, 'pdf')
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        # ✅ Security: Validate file size
        is_valid, error = validate_file_size(file, MAX_FILE_SIZE['pdf'])
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        
        usage_count = get_usage_count('pdf_to_image', ip_address)
        
        if not is_premium and usage_count >= 3:
            return jsonify({
                'success': False,
                'error': f'Free limit reached. You can convert {3 - usage_count} more pages today.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 3,
                'remaining_pages': max(0, 3 - usage_count)
            }), 403
        
        # Save uploaded file with secure filename
        filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
        file.save(temp_path)
        
        # ✅ Security: Check for malicious content
        with open(temp_path, 'rb') as f:
            file_bytes = f.read()
            is_malicious, error = detect_malicious_content(file_bytes)
            if is_malicious:
                cleanup_temp_files([temp_path])
                return jsonify({'success': False, 'error': error}), 400
        
        image_list = []
        total_pages = 0
        
        # ✅ PRIMARY METHOD: PyMuPDF (fitz) - Works on Windows!
        try:
            import fitz
            doc = fitz.open(temp_path)
            total_pages = len(doc)
            limit = total_pages if is_premium else min(3, total_pages)
            
            for i in range(limit):
                page = doc.load_page(i)
                # Higher DPI for better quality
                pix = page.get_pixmap(dpi=200)
                img_data = pix.tobytes('png')
                img_base64 = base64.b64encode(img_data).decode()
                
                # Add watermark for free users
                if should_add_watermark(is_premium):
                    watermarked_img = Image.open(BytesIO(img_data))
                    watermarked_img = add_watermark_to_image(watermarked_img, is_premium)
                    watermarked_buffer = BytesIO()
                    watermarked_img.save(watermarked_buffer, format='PNG')
                    img_base64 = base64.b64encode(watermarked_buffer.getvalue()).decode()
                
                image_list.append({
                    'page': i + 1,
                    'image': f"data:image/png;base64,{img_base64}"
                })
            
            doc.close()
            
        except ImportError:
            # ✅ FALLBACK: Your existing PyPDF2 code (text-based preview)
            with open(temp_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                total_pages = len(pdf_reader.pages)
                limit = total_pages if is_premium else min(3, total_pages)
                
                for i in range(limit):
                    page = pdf_reader.pages[i]
                    text = page.extract_text() or ''
                    
                    img = Image.new('RGB', (800, 1000), color='white')
                    draw = ImageDraw.Draw(img)
                    
                    try:
                        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
                    except:
                        font = ImageFont.load_default()
                    
                    draw.text((50, 50), f"Page {i+1} of {total_pages}", fill='black', font=font)
                    draw.text((50, 80), f"Content preview:", fill='black', font=font)
                    
                    preview_text = text[:200] if text else 'No text extracted (scanned page)'
                    y_pos = 110
                    for line in preview_text.split('\n')[:10]:
                        draw.text((50, y_pos), line[:80], fill='gray', font=font)
                        y_pos += 20
                    
                    img_buffer = BytesIO()
                    img.save(img_buffer, format='PNG')
                    
                    if should_add_watermark(is_premium):
                        watermarked_img = Image.open(img_buffer)
                        watermarked_img = add_watermark_to_image(watermarked_img, is_premium)
                        img_buffer = BytesIO()
                        watermarked_img.save(img_buffer, format='PNG')
                    
                    img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
                    image_list.append({
                        'page': i + 1,
                        'image': f"data:image/png;base64,{img_base64}"
                    })
        
        # Clean up
        cleanup_temp_files([temp_path])
        
        # Increment usage for each page converted
        for _ in range(len(image_list)):
            increment_usage('pdf_to_image', ip_address)
        
        return jsonify({
            'success': True,
            'images': image_list,
            'total_pages': total_pages,
            'converted': len(image_list),
            'usage_count': get_usage_count('pdf_to_image', ip_address),
            'remaining_free': max(0, 3 - get_usage_count('pdf_to_image', ip_address)) if not is_premium else "Unlimited",
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"PDF to Image Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files([temp_path] if 'temp_path' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 5. PDF TO WORD - CLEAN & OPTIMIZED
# ============================================

@tools_bp.route('/pdf-to-word', methods=['POST', 'OPTIONS'])
def pdf_to_word():
    """Convert PDF to Word with priority-based extraction"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    try:
        file = request.files['file']
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'pdf_to_word', limit=15, window=60):
            return jsonify({'success': False, 'error': 'Too many requests. Please try again later.'}), 429
        
        track_request(ip_address, 'pdf_to_word')
        
        # Validate file
        is_valid, error = validate_file_type(file, 'pdf')
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        is_valid, error = validate_file_size(file, MAX_FILE_SIZE['pdf'])
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        # Check limits
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        usage_count = get_usage_count('pdf_to_word', ip_address)
        
        if not is_premium and usage_count >= 2:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 2
            }), 403
        
        # Get options
        options = {}
        try:
            options = json.loads(request.form.get('options', '{}'))
        except:
            pass
        
        include_images = options.get('includeImages', True)
        
        # Save file
        filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
        file.save(temp_path)
        
        # Security check
        with open(temp_path, 'rb') as f:
            file_bytes = f.read()
            is_malicious, error = detect_malicious_content(file_bytes)
            if is_malicious:
                cleanup_temp_files([temp_path])
                return jsonify({'success': False, 'error': error}), 400
        
        doc_data = None
        method_used = "fallback"
        
        # ============================================
        # ✅ PRIORITY 1: pdf2docx (BEST overall)
        # ============================================
        try:
            from pdf2docx import Converter
            output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.docx")
            
            cv = Converter(temp_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            with open(output_path, 'rb') as f:
                doc_data = f.read()
            
            method_used = "pdf2docx"
            cleanup_temp_files([output_path])
            
        except ImportError:
            # ============================================
            # ✅ PRIORITY 2: OCRmyPDF (for scanned PDFs)
            # ============================================
            try:
                import subprocess
                import ocrmypdf
                
                # Check if PDF has text
                has_text = False
                try:
                    import fitz
                    pdf_doc = fitz.open(temp_path)
                    for page in pdf_doc:
                        if page.get_text().strip():
                            has_text = True
                            break
                    pdf_doc.close()
                except:
                    pass
                
                if not has_text:
                    # Apply OCR
                    ocr_output = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.pdf")
                    ocrmypdf.ocr(temp_path, ocr_output, language='eng', force_ocr=True)
                    
                    # Now convert OCR'd PDF with pdf2docx
                    from pdf2docx import Converter
                    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.docx")
                    
                    cv = Converter(ocr_output)
                    cv.convert(output_path, start=0, end=None)
                    cv.close()
                    
                    with open(output_path, 'rb') as f:
                        doc_data = f.read()
                    
                    method_used = "ocr_pdf2docx"
                    cleanup_temp_files([ocr_output, output_path])
                else:
                    # PDF has text but pdf2docx failed, try again with different settings
                    from pdf2docx import Converter
                    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.docx")
                    
                    cv = Converter(temp_path)
                    cv.convert(output_path, start=0, end=None, multi_processing=True)
                    cv.close()
                    
                    with open(output_path, 'rb') as f:
                        doc_data = f.read()
                    
                    method_used = "pdf2docx_retry"
                    cleanup_temp_files([output_path])
                
            except (ImportError, Exception):
                # ============================================
                # ✅ PRIORITY 3: PyMuPDF (good for images)
                # ============================================
                try:
                    import fitz
                    from docx import Document
                    from docx.shared import Inches, Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    
                    doc = Document()
                    title = doc.add_heading('PDF Content', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    pdf_doc = fitz.open(temp_path)
                    
                    # Check if PDF has text
                    has_text = False
                    for page in pdf_doc:
                        if page.get_text().strip():
                            has_text = True
                            break
                    
                    if not has_text:
                        # Image PDF - extract images
                        doc.add_heading('Image PDF - Extracted as Images', level=1)
                        for page_num, page in enumerate(pdf_doc):
                            doc.add_heading(f'Page {page_num + 1}', level=2)
                            pix = page.get_pixmap(dpi=150)
                            img_data = pix.tobytes("png")
                            img_buffer = BytesIO(img_data)
                            doc.add_picture(img_buffer, width=Inches(5))
                            doc.add_paragraph()
                        method_used = "pymupdf_images"
                    else:
                        # Text PDF - extract content
                        for page_num, page in enumerate(pdf_doc):
                            doc.add_heading(f'Page {page_num + 1}', level=1)
                            
                            # Extract images
                            if include_images:
                                image_list = page.get_images(full=True)
                                for img in image_list:
                                    try:
                                        xref = img[0]
                                        pix = fitz.Pixmap(pdf_doc, xref)
                                        if pix.n - pix.alpha < 4:
                                            img_data = pix.tobytes("png")
                                            img_buffer = BytesIO(img_data)
                                            doc.add_picture(img_buffer, width=Inches(4))
                                            doc.add_paragraph()
                                        pix = None
                                    except:
                                        pass
                            
                            # Extract tables
                            tables = page.find_tables()
                            if tables:
                                for table in tables:
                                    table_data = table.extract()
                                    if table_data and len(table_data) > 1:
                                        max_cols = max(len(row) for row in table_data)
                                        word_table = doc.add_table(rows=len(table_data), cols=max_cols)
                                        word_table.style = 'Table Grid'
                                        
                                        for row_idx, row_data in enumerate(table_data):
                                            for col_idx in range(max_cols):
                                                cell_text = str(row_data[col_idx]) if col_idx < len(row_data) and row_data[col_idx] is not None else ''
                                                word_table.cell(row_idx, col_idx).text = cell_text.strip()
                                                
                                                if row_idx == 0:
                                                    for paragraph in word_table.cell(row_idx, col_idx).paragraphs:
                                                        for run in paragraph.runs:
                                                            run.font.bold = True
                                                            run.font.size = Pt(11)
                                        doc.add_paragraph()
                            
                            # Extract text
                            text = page.get_text()
                            if text:
                                lines = text.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if line:
                                        if line.isupper() and len(line) > 5:
                                            doc.add_heading(line, level=2)
                                        else:
                                            doc.add_paragraph(line)
                    
                    pdf_doc.close()
                    doc_bytes = BytesIO()
                    doc.save(doc_bytes)
                    doc_data = doc_bytes.getvalue()
                    if method_used != "pymupdf_images":
                        method_used = "pymupdf"
                    
                except ImportError:
                    # ============================================
                    # ✅ PRIORITY 4: pdfplumber (good for tables)
                    # ============================================
                    try:
                        import pdfplumber
                        from docx import Document
                        from docx.shared import Inches, Pt
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.enum.table import WD_TABLE_ALIGNMENT
                        import re
                        
                        doc = Document()
                        title = doc.add_heading('PDF Content', 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        with pdfplumber.open(temp_path) as pdf:
                            # Check if PDF has text
                            has_text = False
                            for page in pdf.pages:
                                if page.extract_text():
                                    has_text = True
                                    break
                            
                            if not has_text:
                                doc.add_paragraph("(Image PDF - No text extracted)")
                                method_used = "pdfplumber_no_text"
                            else:
                                for page_num, page in enumerate(pdf.pages):
                                    doc.add_heading(f'Page {page_num + 1}', level=1)
                                    
                                    # Extract tables
                                    tables = page.extract_tables()
                                    unique_tables = []
                                    seen = set()
                                    
                                    for table_data in tables:
                                        if table_data and len(table_data) > 1:
                                            table_hash = str(table_data)
                                            if table_hash not in seen:
                                                seen.add(table_hash)
                                                unique_tables.append(table_data)
                                    
                                    for table_data in unique_tables:
                                        filtered_rows = []
                                        for row in table_data:
                                            if any(cell and str(cell).strip() for cell in row):
                                                filtered_rows.append(row)
                                        
                                        if filtered_rows:
                                            max_cols = max(len(row) for row in filtered_rows)
                                            table = doc.add_table(rows=len(filtered_rows), cols=max_cols)
                                            table.style = 'Table Grid'
                                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                            
                                            for row_idx, row_data in enumerate(filtered_rows):
                                                for col_idx in range(max_cols):
                                                    cell_text = str(row_data[col_idx]) if col_idx < len(row_data) and row_data[col_idx] is not None else ''
                                                    table.cell(row_idx, col_idx).text = cell_text.strip()
                                                    
                                                    if row_idx == 0:
                                                        for paragraph in table.cell(row_idx, col_idx).paragraphs:
                                                            for run in paragraph.runs:
                                                                run.font.bold = True
                                                                run.font.size = Pt(11)
                                            doc.add_paragraph()
                                    
                                    # Extract text
                                    text = page.extract_text()
                                    if text:
                                        lines = text.split('\n')
                                        for line in lines:
                                            line = line.strip()
                                            if line:
                                                if line.isupper() and len(line) > 5:
                                                    doc.add_heading(line, level=2)
                                                else:
                                                    doc.add_paragraph(line)
                                
                                method_used = "pdfplumber"
                        
                        doc_bytes = BytesIO()
                        doc.save(doc_bytes)
                        doc_data = doc_bytes.getvalue()
                        
                    except ImportError:
                        # ============================================
                        # ✅ PRIORITY 5: PyPDF2 (emergency fallback)
                        # ============================================
                        from docx import Document
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        
                        doc = Document()
                        title = doc.add_heading('PDF Content', 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        with open(temp_path, 'rb') as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                            for page_num, page in enumerate(pdf_reader.pages):
                                doc.add_heading(f'Page {page_num + 1}', level=1)
                                text = page.extract_text() or ''
                                
                                if not text:
                                    doc.add_paragraph("(Image PDF - No text extracted)")
                                else:
                                    for line in text.split('\n'):
                                        line = line.strip()
                                        if line:
                                            if line.isupper() and len(line) > 5:
                                                doc.add_heading(line, level=2)
                                            else:
                                                doc.add_paragraph(line)
                        
                        doc_bytes = BytesIO()
                        doc.save(doc_bytes)
                        doc_data = doc_bytes.getvalue()
                        method_used = "pypdf2_emergency"
        
        # ✅ Add watermark for free users
        if should_add_watermark(is_premium) and doc_data:
            from docx import Document
            temp_doc = Document(BytesIO(doc_data))
            temp_doc.add_paragraph()
            temp_doc.add_paragraph('─' * 50)
            temp_doc.add_paragraph('Made with ❤️ by Krynova Technologies')
            temp_doc.add_paragraph('Visit: https://krynovatechnology.pythonanywhere.com')
            temp_doc.add_paragraph('─' * 50)
            watermark_bytes = BytesIO()
            temp_doc.save(watermark_bytes)
            doc_data = watermark_bytes.getvalue()
        
        # Cleanup and response
        cleanup_temp_files([temp_path])
        increment_usage('pdf_to_word', ip_address)
        
        return jsonify({
            'success': True,
            'file': base64.b64encode(doc_data).decode() if doc_data else '',
            'filename': filename.replace('.pdf', '.docx'),
            'usage_count': get_usage_count('pdf_to_word', ip_address),
            'remaining_free': max(0, 2 - get_usage_count('pdf_to_word', ip_address)) if not is_premium else "Unlimited",
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium),
            'method_used': method_used,
            'message': f'Converted using {method_used}'
        })
        
    except Exception as e:
        print(f"PDF to Word Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files([temp_path] if 'temp_path' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 6. PDF TO EXCEL - ENHANCED
# ============================================

@tools_bp.route('/pdf-to-excel', methods=['POST', 'OPTIONS'])
def pdf_to_excel():
    """Convert PDF to Excel with tables, images, and formatting preservation"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    try:
        file = request.files['file']
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'pdf_to_excel', limit=15, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'pdf_to_excel')
        
        # ✅ Security: Validate file type
        is_valid, error = validate_file_type(file, 'pdf')
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        # ✅ Security: Validate file size
        is_valid, error = validate_file_size(file, MAX_FILE_SIZE['pdf'])
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        user_id = request.form.get('user_id', 'anonymous')
        
        usage_count = get_usage_count('pdf_to_excel', ip_address)
        
        if not is_premium and usage_count >= 2:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 2
            }), 403
        
        filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
        file.save(temp_path)
        
        # ✅ Security: Check for malicious content
        with open(temp_path, 'rb') as f:
            file_bytes = f.read()
            is_malicious, error = detect_malicious_content(file_bytes)
            if is_malicious:
                cleanup_temp_files([temp_path])
                return jsonify({'success': False, 'error': error}), 400
        
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        # ✅ Track sheets to avoid duplicate names
        sheet_names = set()
        
        # ✅ METHOD 1: Try tabula-py (best table extraction)
        try:
            import tabula
            import pandas as pd
            
            tables = tabula.read_pdf(temp_path, pages='all', multiple_tables=True)
            
            if tables:
                for i, table in enumerate(tables):
                    if not table.empty:
                        # Clean column names
                        table.columns = [str(col).strip() if pd.notna(col) else f'Column_{j+1}' for j, col in enumerate(table.columns)]
                        
                        sheet_name = f'Table_{i+1}'[:31]
                        # Handle duplicate sheet names
                        if sheet_name in sheet_names:
                            sheet_name = f'Table_{i+1}_v2'[:31]
                        sheet_names.add(sheet_name)
                        
                        ws = wb.create_sheet(title=sheet_name)
                        
                        # ✅ Style the header
                        header_font = Font(bold=True, color='FFFFFF', size=11)
                        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                        header_alignment = Alignment(horizontal='center', vertical='center')
                        
                        # Write headers with styling
                        for col_idx, header in enumerate(table.columns, 1):
                            cell = ws.cell(row=1, column=col_idx, value=str(header))
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = header_alignment
                        
                        # Write data with formatting
                        border = Border(
                            left=Side(style='thin'),
                            right=Side(style='thin'),
                            top=Side(style='thin'),
                            bottom=Side(style='thin')
                        )
                        
                        for row_idx, row in enumerate(table.values, 2):
                            for col_idx, value in enumerate(row, 1):
                                if pd.notna(value):
                                    # Preserve numbers as numbers
                                    if isinstance(value, (int, float)):
                                        cell = ws.cell(row=row_idx, column=col_idx, value=value)
                                    else:
                                        cell = ws.cell(row=row_idx, column=col_idx, value=str(value))
                                    
                                    cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                                    cell.border = border
                        
                        # Auto-fit columns
                        for col_idx in range(1, len(table.columns) + 1):
                            max_length = 0
                            column = get_column_letter(col_idx)
                            for row in range(1, ws.max_row + 1):
                                cell_value = ws.cell(row=row, column=col_idx).value
                                if cell_value:
                                    max_length = max(max_length, len(str(cell_value)))
                            ws.column_dimensions[column].width = min(max_length + 3, 50)
            else:
                # No tables found - extract text with structure
                ws = wb.create_sheet(title="Extracted_Text")
                row_num = 1
                
                # ✅ Use PyMuPDF if available for better text extraction
                try:
                    import fitz
                    doc = fitz.open(temp_path)
                    for page in doc:
                        text = page.get_text()
                        for line in text.split('\n'):
                            if line.strip():
                                # Try to split by multiple spaces (table-like)
                                parts = re.split(r'\s{2,}', line.strip())
                                if len(parts) >= 3:
                                    for col_idx, part in enumerate(parts, 1):
                                        ws.cell(row=row_num, column=col_idx, value=part.strip())
                                else:
                                    ws.cell(row=row_num, column=1, value=line.strip())
                                row_num += 1
                    doc.close()
                except:
                    # Fallback to PyPDF2
                    with open(temp_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text = page.extract_text() or ''
                            for line in text.split('\n'):
                                if line.strip():
                                    # Try to split by multiple spaces (table-like)
                                    parts = re.split(r'\s{2,}', line.strip())
                                    if len(parts) >= 3:
                                        for col_idx, part in enumerate(parts, 1):
                                            ws.cell(row=row_num, column=col_idx, value=part.strip())
                                    else:
                                        ws.cell(row=row_num, column=1, value=line.strip())
                                    row_num += 1
        
        except ImportError:
            # ✅ METHOD 2: Fallback - Use PyMuPDF for better extraction
            ws = wb.create_sheet(title="PDF_Content")
            row_num = 1
            
            try:
                import fitz
                import re
                
                doc = fitz.open(temp_path)
                
                for page_num, page in enumerate(doc):
                    # Add page header
                    ws.cell(row=row_num, column=1, value=f"=== Page {page_num + 1} ===")
                    row_num += 1
                    
                    # Extract text with position info
                    blocks = page.get_text("dict")
                    
                    for block in blocks.get("blocks", []):
                        if block.get("type") == 0:  # Text block
                            for line in block.get("lines", []):
                                spans = line.get("spans", [])
                                if spans:
                                    text = "".join([span.get("text", "") for span in spans])
                                    if text.strip():
                                        # Try to detect table rows
                                        parts = re.split(r'\s{2,}', text.strip())
                                        if len(parts) >= 3:
                                            for col_idx, part in enumerate(parts, 1):
                                                ws.cell(row=row_num, column=col_idx, value=part.strip())
                                        else:
                                            ws.cell(row=row_num, column=1, value=text.strip())
                                        row_num += 1
                
                doc.close()
                
            except:
                # ✅ METHOD 3: Final fallback - PyPDF2
                with open(temp_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text = page.extract_text() or ''
                        for line in text.split('\n'):
                            if line.strip():
                                ws.cell(row=row_num, column=1, value=line.strip())
                                row_num += 1
        
        # ✅ Add summary sheet
        summary_ws = wb.create_sheet(title="Summary")
        summary_ws.cell(row=1, column=1, value="PDF to Excel Conversion Summary")
        summary_ws.cell(row=1, column=1).font = Font(bold=True, size=14)
        
        summary_data = [
            ["Original File", filename],
            ["Conversion Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ["Total Sheets", len(wb.sheetnames)],
            ["Premium User", "Yes" if is_premium else "No"],
            ["Free Uses Left", str(max(0, 2 - usage_count)) if not is_premium else "Unlimited"],
        ]
        
        for row_idx, (key, value) in enumerate(summary_data, 3):
            summary_ws.cell(row=row_idx, column=1, value=key).font = Font(bold=True)
            summary_ws.cell(row=row_idx, column=2, value=value)
        
        # Add watermark for free users
        if should_add_watermark(is_premium):
            # Add watermark text to all sheets
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                if ws.title != "Summary":
                    max_row = ws.max_row + 2
                    ws.cell(row=max_row, column=1, value='─' * 30)
                    ws.cell(row=max_row+1, column=1, value='Made with ❤️ by Krynova Technologies')
                    ws.cell(row=max_row+2, column=1, value='Visit: https://krynovatechnology.pythonanywhere.com')
                    ws.cell(row=max_row+3, column=1, value='─' * 30)
        
        excel_bytes = BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        excel_base64 = base64.b64encode(excel_bytes.read()).decode()
        
        cleanup_temp_files([temp_path])
        increment_usage('pdf_to_excel', ip_address)
        
        return jsonify({
            'success': True,
            'file': excel_base64,
            'filename': filename.replace('.pdf', '.xlsx'),
            'sheets': wb.sheetnames,
            'sheet_count': len(wb.sheetnames),
            'usage_count': get_usage_count('pdf_to_excel', ip_address),
            'remaining_free': max(0, 2 - get_usage_count('pdf_to_excel', ip_address)) if not is_premium else "Unlimited",
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium),
            'message': 'Conversion successful' if is_premium else 'Free conversion - Upgrade for unlimited'
        })
        
    except Exception as e:
        print(f"PDF to Excel Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files([temp_path] if 'temp_path' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 7. IMAGE TO PDF
# ============================================

@tools_bp.route('/image-to-pdf', methods=['POST', 'OPTIONS'])
def image_to_pdf():
    """Convert Images to PDF (Free: Single images unlimited, Bulk: 2 per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'files' not in request.files:
        return jsonify({'success': False, 'error': 'No files uploaded'}), 400
    
    try:
        files = request.files.getlist('files')
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'image_to_pdf', limit=15, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'image_to_pdf')
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        
        # ✅ Security: Validate each file
        for file in files:
            is_valid, error = validate_file_type(file, 'image')
            if not is_valid:
                return jsonify({'success': False, 'error': f'Invalid file: {file.filename} - {error}'}), 400
            
            is_valid, error = validate_file_size(file, MAX_FILE_SIZE['image'])
            if not is_valid:
                return jsonify({'success': False, 'error': f'File too large: {file.filename} - {error}'}), 400
        
        # Get options from request
        options = {}
        try:
            options = json.loads(request.form.get('options', '{}'))
        except:
            pass
        
        # ✅ FIX: Single image = unlimited, Bulk = limited
        is_bulk = len(files) > 1
        
        # Track bulk usage for free users
        bulk_key = f"bulk_image_to_pdf:{ip_address}"
        today = datetime.now().strftime('%Y-%m-%d')
        
        if not is_premium and is_bulk:
            bulk_data = usage_tracking.get(bulk_key, {'date': '', 'count': 0})
            if bulk_data['date'] != today:
                bulk_data = {'date': today, 'count': 0}
            
            if bulk_data['count'] >= 2:
                return jsonify({
                    'success': False,
                    'error': 'Bulk conversion limit reached. Free users can convert 2 bulk (2+ images) per day. Upgrade to premium for unlimited.',
                    'limit_reached': True,
                    'limit_type': 'bulk',
                    'max_free': 2,
                    'used_today': bulk_data['count']
                }), 403
            
            bulk_data['count'] += 1
            usage_tracking[bulk_key] = bulk_data
        
        # Process images to PDF using PIL
        image_list = []
        temp_paths = []
        
        for file in files:
            filename = sanitize_filename(file.filename)
            temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
            file.save(temp_path)
            temp_paths.append(temp_path)
            
            # Open and convert image
            img = Image.open(temp_path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize if too large (max 2000px)
            max_dim = 2000
            if img.width > max_dim or img.height > max_dim:
                img.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS)
            
            image_list.append(img)
        
        # Create PDF from images
        pdf_buffer = BytesIO()
        if image_list:
            image_list[0].save(
                pdf_buffer,
                format='PDF',
                save_all=True,
                append_images=image_list[1:] if len(image_list) > 1 else [],
                resolution=100.0,
                quality=85 if not is_premium else 95
            )
        
        pdf_bytes = pdf_buffer.getvalue()
        
        # ✅ Add watermark for free users (as text footer using PyPDF2)
        if should_add_watermark(is_premium):
            try:
                from PyPDF2 import PdfReader, PdfWriter
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                from io import BytesIO
                
                # Create watermark page
                watermark_buffer = BytesIO()
                c = canvas.Canvas(watermark_buffer, pagesize=letter)
                c.setFont("Helvetica", 10)
                c.setFillColorRGB(0.5, 0.5, 0.5, 0.3)
                c.drawRightString(500, 30, "Made with ❤️ by Krynova")
                c.save()
                watermark_buffer.seek(0)
                
                # Merge with PDF
                reader = PdfReader(BytesIO(pdf_bytes))
                writer = PdfWriter()
                watermark_reader = PdfReader(watermark_buffer)
                watermark_page = watermark_reader.pages[0]
                
                for page in reader.pages:
                    page.merge_page(watermark_page)
                    writer.add_page(page)
                
                final_buffer = BytesIO()
                writer.write(final_buffer)
                pdf_bytes = final_buffer.getvalue()
            except Exception as e:
                print(f"Watermark error: {str(e)}")
        
        pdf_base64 = base64.b64encode(pdf_bytes).decode()
        
        # Clean up
        cleanup_temp_files(temp_paths)
        
        # For single images, only track total usage for display purposes
        if not is_premium:
            increment_usage('image_to_pdf', ip_address)
            current_usage = get_usage_count('image_to_pdf', ip_address)
            remaining = max(0, 999 - current_usage)
        else:
            current_usage = 0
            remaining = "Unlimited"
        
        return jsonify({
            'success': True,
            'file': pdf_base64,
            'filename': f'converted_{datetime.now().strftime("%Y%m%d")}.pdf',
            'pages': len(image_list),
            'usage_count': current_usage,
            'remaining_free': remaining,
            'is_premium': is_premium,
            'is_bulk': is_bulk,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"Image to PDF Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files(temp_paths if 'temp_paths' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 8. PDF COMPRESSOR - ENHANCED
# ============================================

@tools_bp.route('/pdf-compressor', methods=['POST', 'OPTIONS'])
def compress_pdf():
    """Compress PDF with multiple methods for maximum size reduction"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    try:
        file = request.files['file']
        ip_address = get_client_ip(request)
        
        if is_rate_limited(ip_address, 'pdf_compressor', limit=15, window=60):
            return jsonify({'success': False, 'error': 'Too many requests. Please try again later.'}), 429
        
        track_request(ip_address, 'pdf_compressor')
        
        is_valid, error = validate_file_type(file, 'pdf')
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        is_valid, error = validate_file_size(file, MAX_FILE_SIZE['pdf'])
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        user_id = request.form.get('user_id', 'anonymous')
        
        options = {}
        try:
            options = json.loads(request.form.get('options', '{}'))
        except:
            pass
        
        target_size_enabled = options.get('targetSize', False)
        target_size_kb = options.get('targetSizeKB', 20)
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        min_target_size = 5 if is_premium else 20
        
        if target_size_enabled:
            if target_size_kb < min_target_size:
                return jsonify({
                    'success': False,
                    'error': f'Target size must be at least {min_target_size}KB. {"Upgrade to premium for 5KB target." if not is_premium else ""}',
                    'min_target_size': min_target_size,
                    'is_premium': is_premium
                }), 400
            if target_size_kb > 10000:
                return jsonify({
                    'success': False,
                    'error': 'Target size cannot exceed 10MB (10000KB)'
                }), 400
        
        is_batch = request.form.get('batch_mode', 'false').lower() == 'true'
        
        batch_key = f"batch_pdf_compressor:{ip_address}"
        today = datetime.now().strftime('%Y-%m-%d')
        
        if not is_premium and is_batch:
            batch_data = usage_tracking.get(batch_key, {'date': '', 'count': 0})
            if batch_data['date'] != today:
                batch_data = {'date': today, 'count': 0}
            
            if batch_data['count'] >= 3:
                return jsonify({
                    'success': False,
                    'error': 'Batch compression limit reached. Free users can compress 3 batches per day. Upgrade to premium for unlimited.',
                    'limit_reached': True,
                    'limit_type': 'batch',
                    'max_free': 3,
                    'used_today': batch_data['count']
                }), 403
            
            batch_data['count'] += 1
            usage_tracking[batch_key] = batch_data
        
        filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
        file.save(temp_path)
        
        with open(temp_path, 'rb') as f:
            file_bytes = f.read()
            is_malicious, error = detect_malicious_content(file_bytes)
            if is_malicious:
                cleanup_temp_files([temp_path])
                return jsonify({'success': False, 'error': error}), 400
        
        original_size = os.path.getsize(temp_path)
        compressed_data = None
        compressed_size = original_size
        saved_percentage = 0
        target_size_reached = False
        
        # ✅ METHOD 1: PyPDF2 compression (works for text-based PDFs)
        try:
            with open(temp_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                pdf_writer = PyPDF2.PdfWriter()
                
                for page in pdf_reader.pages:
                    try:
                        page.compress_content_streams()
                    except:
                        pass
                    pdf_writer.add_page(page)
                
                pdf_writer.add_metadata({
                    '/Creator': 'Krynova PDF Compressor',
                    '/Producer': 'Krynova'
                })
                
                compressed_buffer = BytesIO()
                pdf_writer.write(compressed_buffer)
                test_data = compressed_buffer.getvalue()
                
                if len(test_data) < original_size and len(test_data) > 0:
                    compressed_data = test_data
                    compressed_size = len(test_data)
                    
                    if target_size_enabled and compressed_size <= (target_size_kb * 1024):
                        target_size_reached = True
        except:
            pass
        
        # ✅ METHOD 2: Ghostscript compression (best for images - if available)
        if (not compressed_data or compressed_size >= original_size) and not target_size_reached:
            try:
                import subprocess
                output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.pdf")
                
                # Use different settings for better compression
                if target_size_enabled:
                    if target_size_kb <= 10:
                        settings = "screen"
                    elif target_size_kb <= 50:
                        settings = "ebook"
                    else:
                        settings = "printer" if is_premium else "ebook"
                else:
                    settings = "ebook" if not is_premium else "printer"
                
                # Try multiple passes for better compression
                for quality in ['screen', 'ebook', 'printer']:
                    output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.pdf")
                    
                    cmd = [
                        'gs', '-sDEVICE=pdfwrite',
                        '-dCompatibilityLevel=1.4',
                        f'-dPDFSETTINGS=/{quality}',
                        '-dNOPAUSE', '-dQUIET', '-dBATCH',
                        '-dDetectDuplicateImages=true',
                        '-dCompressFonts=true',
                        '-dSubsetFonts=true',
                        '-dEmbedAllFonts=false',
                        '-dMaxSubsetPct=100',
                        '-dDownsampleColorImages=true',
                        '-dDownsampleGrayImages=true',
                        '-dDownsampleMonoImages=true',
                        '-dColorImageResolution=150',
                        '-dGrayImageResolution=150',
                        '-dMonoImageResolution=150',
                        f'-sOutputFile={output_path}',
                        temp_path
                    ]
                    
                    try:
                        subprocess.run(cmd, capture_output=True, timeout=30)
                        if os.path.exists(output_path):
                            with open(output_path, 'rb') as f:
                                test_data = f.read()
                            if len(test_data) < original_size and len(test_data) > 0:
                                compressed_data = test_data
                                compressed_size = len(test_data)
                                
                                if target_size_enabled and compressed_size <= (target_size_kb * 1024):
                                    target_size_reached = True
                                    break
                            cleanup_temp_files([output_path])
                    except:
                        pass
            except:
                pass
        
        # ✅ METHOD 3: img2pdf (for image-heavy PDFs)
        if (not compressed_data or compressed_size >= original_size) and not target_size_reached:
            try:
                import img2pdf
                from PIL import Image
                
                # Convert PDF to images and back to PDF with compression
                try:
                    import fitz
                    doc = fitz.open(temp_path)
                    images = []
                    
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(dpi=100)  # Lower DPI for compression
                        img_data = pix.tobytes("png")
                        images.append(Image.open(BytesIO(img_data)))
                    
                    doc.close()
                    
                    if images:
                        pdf_buffer = BytesIO()
                        images[0].save(
                            pdf_buffer,
                            format='PDF',
                            save_all=True,
                            append_images=images[1:] if len(images) > 1 else [],
                            quality=70 if not is_premium else 85,
                            optimize=True
                        )
                        test_data = pdf_buffer.getvalue()
                        
                        if len(test_data) < original_size and len(test_data) > 0:
                            compressed_data = test_data
                            compressed_size = len(test_data)
                            
                            if target_size_enabled and compressed_size <= (target_size_kb * 1024):
                                target_size_reached = True
                except:
                    pass
            except:
                pass
        
        # ✅ METHOD 4: pypdf2 with aggressive settings
        if (not compressed_data or compressed_size >= original_size) and not target_size_reached:
            try:
                from PyPDF2 import PdfReader, PdfWriter
                
                reader = PdfReader(temp_path)
                writer = PdfWriter()
                
                for page in reader.pages:
                    # Compress content streams
                    try:
                        page.compress_content_streams()
                    except:
                        pass
                    writer.add_page(page)
                
                # Remove all metadata
                writer.add_metadata({})
                
                # Compress with different settings
                compressed_buffer = BytesIO()
                writer.write(compressed_buffer)
                test_data = compressed_buffer.getvalue()
                
                if len(test_data) < original_size and len(test_data) > 0:
                    compressed_data = test_data
                    compressed_size = len(test_data)
                    
                    if target_size_enabled and compressed_size <= (target_size_kb * 1024):
                        target_size_reached = True
            except:
                pass
        
        # If still no compression, use original
        if not compressed_data or compressed_size >= original_size:
            with open(temp_path, 'rb') as f:
                compressed_data = f.read()
            compressed_size = original_size
            saved_percentage = 0
        else:
            saved_percentage = int((1 - compressed_size / original_size) * 100)
            if saved_percentage < 0:
                saved_percentage = 0
        
        # Clean up
        cleanup_temp_files([temp_path])
        
        usage_count = get_usage_count('pdf_compressor', ip_address)
        
        if not is_premium:
            increment_usage('pdf_compressor', ip_address)
            current_usage = get_usage_count('pdf_compressor', ip_address)
            if is_batch:
                remaining = max(0, 3 - batch_data['count'])
            else:
                remaining = "Unlimited"
        else:
            current_usage = 0
            remaining = "Unlimited"
        
        pdf_base64 = base64.b64encode(compressed_data).decode()
        
        # ✅ If target size not reached, provide feedback
        compression_message = f"Saved {saved_percentage}%"
        if target_size_enabled and not target_size_reached:
            compression_message = f"Saved {saved_percentage}% (Target {target_size_kb}KB not reached. Try using 'Maximum Size Reduction' mode.)"
        elif target_size_enabled and target_size_reached:
            compression_message = f"✅ Target size {target_size_kb}KB reached! Saved {saved_percentage}%"
        
        return jsonify({
            'success': True,
            'file': pdf_base64,
            'filename': f'compressed_{filename}',
            'original_size': original_size,
            'compressed_size': compressed_size,
            'saved_percentage': saved_percentage,
            'usage_count': current_usage,
            'remaining_free': remaining,
            'is_premium': is_premium,
            'is_batch': is_batch,
            'batch_remaining': max(0, 3 - batch_data['count']) if not is_premium and is_batch else "Unlimited",
            'has_watermark': False,
            'target_size_reached': target_size_reached,
            'target_size_kb': target_size_kb if target_size_enabled else None,
            'min_target_size': min_target_size,
            'user_id': user_id,
            'compression_message': compression_message
        })
        
    except Exception as e:
        print(f"PDF Compressor Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files([temp_path] if 'temp_path' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 9. MERGE PDF
# ============================================

@tools_bp.route('/merge-pdf', methods=['POST', 'OPTIONS'])
def merge_pdf():
    """Merge PDFs (Free: Up to 35 files, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'files' not in request.files:
        return jsonify({'success': False, 'error': 'No files uploaded'}), 400
    
    try:
        files = request.files.getlist('files')
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'merge_pdf', limit=10, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'merge_pdf')
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        
        # ✅ Security: Validate each file
        for file in files:
            is_valid, error = validate_file_type(file, 'pdf')
            if not is_valid:
                return jsonify({'success': False, 'error': f'Invalid file: {file.filename} - {error}'}), 400
            
            is_valid, error = validate_file_size(file, MAX_FILE_SIZE['pdf'])
            if not is_valid:
                return jsonify({'success': False, 'error': f'File too large: {file.filename} - {error}'}), 400
        
        # ✅ Check file count limit: Free = 35 max, Premium = Unlimited
        MAX_FILES_FREE = 35
        if not is_premium and len(files) > MAX_FILES_FREE:
            return jsonify({
                'success': False,
                'error': f'Free users can merge up to {MAX_FILES_FREE} files. You have {len(files)} files. Please upgrade to premium for unlimited files.',
                'limit_reached': True,
                'limit_type': 'file_count',
                'max_free': MAX_FILES_FREE,
                'file_count': len(files)
            }), 403
        
        # Check daily usage limit (3 per day for free)
        usage_count = get_usage_count('merge_pdf', ip_address)
        
        if not is_premium and usage_count >= 3:
            return jsonify({
                'success': False,
                'error': 'Daily merge limit reached (3 per day). Please upgrade to premium for unlimited merges.',
                'limit_reached': True,
                'limit_type': 'daily',
                'usage_count': usage_count,
                'max_free': 3
            }), 403
        
        # Get options
        options = {}
        try:
            options = json.loads(request.form.get('options', '{}'))
        except:
            pass
        
        # Create merger
        merger = PyPDF2.PdfMerger()
        temp_paths = []
        file_names = []
        
        # Process files one by one
        for idx, file in enumerate(files):
            filename = sanitize_filename(file.filename)
            file_names.append(filename)
            temp_path = os.path.join(tempfile.gettempdir(), f"merge_{uuid.uuid4().hex}_{filename}")
            file.save(temp_path)
            temp_paths.append(temp_path)
            
            # ✅ Security: Check for malicious content
            with open(temp_path, 'rb') as f:
                file_bytes = f.read()
                is_malicious, error = detect_malicious_content(file_bytes)
                if is_malicious:
                    cleanup_temp_files(temp_paths)
                    return jsonify({'success': False, 'error': error}), 400
            
            try:
                with open(temp_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    page_count = len(reader.pages)
                merger.append(temp_path)
            except Exception as e:
                cleanup_temp_files(temp_paths)
                return jsonify({
                    'success': False,
                    'error': f'Error processing file "{filename}". Please check if it\'s a valid PDF.'
                }), 400
        
        # Write merged PDF
        output_bytes = BytesIO()
        merger.write(output_bytes)
        merger.close()
        
        merged_data = output_bytes.getvalue()
        
        # Count total pages
        try:
            with BytesIO(merged_data) as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
        except:
            total_pages = len(files)
        
        # Clean up temp files
        cleanup_temp_files(temp_paths)
        
        # Only increment usage for free users
        if not is_premium:
            increment_usage('merge_pdf', ip_address)
            remaining_free = max(0, 3 - get_usage_count('merge_pdf', ip_address))
        else:
            remaining_free = "Unlimited"
        
        pdf_base64 = base64.b64encode(merged_data).decode()
        
        return jsonify({
            'success': True,
            'file': pdf_base64,
            'filename': f'merged_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
            'pages': len(files),
            'total_pages': total_pages,
            'file_count': len(files),
            'file_names': file_names,
            'usage_count': get_usage_count('merge_pdf', ip_address) if not is_premium else 0,
            'remaining_free': remaining_free,
            'max_files_free': MAX_FILES_FREE,
            'is_premium': is_premium,
            'has_watermark': False  # No watermark for merge
        })
        
    except Exception as e:
        print(f"Merge PDF Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files(temp_paths if 'temp_paths' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 10. SPLIT PDF
# ============================================

@tools_bp.route('/split-pdf', methods=['POST', 'OPTIONS'])
def split_pdf():
    """Split PDF (Free: 3 per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    try:
        file = request.files['file']
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'split_pdf', limit=15, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'split_pdf')
        
        # ✅ Security: Validate file type
        is_valid, error = validate_file_type(file, 'pdf')
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        # ✅ Security: Validate file size
        is_valid, error = validate_file_size(file, MAX_FILE_SIZE['pdf'])
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        
        usage_count = get_usage_count('split_pdf', ip_address)
        
        if not is_premium and usage_count >= 3:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 3
            }), 403
        
        filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
        file.save(temp_path)
        
        # ✅ Security: Check for malicious content
        with open(temp_path, 'rb') as f:
            file_bytes = f.read()
            is_malicious, error = detect_malicious_content(file_bytes)
            if is_malicious:
                cleanup_temp_files([temp_path])
                return jsonify({'success': False, 'error': error}), 400
        
        with open(temp_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            
            # Limit to 100 pages to prevent abuse
            max_pages = min(len(reader.pages), 100)
            
            for i in range(max_pages):
                writer = PyPDF2.PdfWriter()
                writer.add_page(reader.pages[i])
                
                page_bytes = BytesIO()
                writer.write(page_bytes)
                pages.append(base64.b64encode(page_bytes.getvalue()).decode())
        
        cleanup_temp_files([temp_path])
        
        increment_usage('split_pdf', ip_address)
        
        return jsonify({
            'success': True,
            'pages': pages,
            'total_pages': len(pages),
            'filename': filename.replace('.pdf', ''),
            'usage_count': get_usage_count('split_pdf', ip_address),
            'remaining_free': max(0, 3 - get_usage_count('split_pdf', ip_address)),
            'is_premium': is_premium,
            'has_watermark': False  # No watermark for split
        })
        
    except Exception as e:
        print(f"Split PDF Error: {str(e)}")
        cleanup_temp_files([temp_path] if 'temp_path' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 11. IMAGE RESIZER
# ============================================

@tools_bp.route('/image-resizer', methods=['POST', 'OPTIONS'])
def resize_image():
    """Resize Image (Free: Unlimited, max 1200x1200, Premium: Unlimited, up to 8000x8000)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    
    try:
        file = request.files['file']
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'image_resizer', limit=30, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'image_resizer')
        
        # ✅ Security: Validate file type
        is_valid, error = validate_file_type(file, 'image')
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        # ✅ Security: Validate file size
        is_valid, error = validate_file_size(file, MAX_FILE_SIZE['image'])
        if not is_valid:
            return jsonify({'success': False, 'error': error}), 400
        
        is_premium = request.form.get('is_premium', 'false').lower() == 'true'
        
        width = int(request.form.get('width', 800))
        height = int(request.form.get('height', 600))
        
        # ✅ Free: Max 1200x1200, Premium: Up to 8000x8000
        max_dimension = 8000 if is_premium else 1200
        
        if width < 10 or height < 10:
            return jsonify({'success': False, 'error': 'Width and height must be at least 10px'}), 400
        
        if width > max_dimension or height > max_dimension:
            return jsonify({
                'success': False, 
                'error': f'Maximum dimension is {max_dimension}px. {"Upgrade to premium for larger sizes." if not is_premium else ""}',
                'limit_reached': True if not is_premium else False,
                'max_dimension': max_dimension,
                'is_premium': is_premium
            }), 400
        
        # Get options
        options = {}
        try:
            options = json.loads(request.form.get('options', '{}'))
        except:
            pass
        
        filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{filename}")
        file.save(temp_path)
        
        # ✅ Open the actual image
        img = Image.open(temp_path)
        original_size = img.size
        
        # ✅ Quality settings
        quality = options.get('quality', 'medium')
        quality_map = {
            'high': 95,
            'medium': 85,
            'low': 70
        }
        quality_value = quality_map.get(quality, 85)
        
        # ✅ Premium gets advanced resizing options
        if is_premium:
            resample = Image.Resampling.LANCZOS
            output_format = options.get('format', 'PNG')
            background = options.get('background', None)
        else:
            resample = Image.Resampling.BICUBIC
            output_format = 'PNG'
            background = None
        
        # ✅ ACTUALLY RESIZE THE IMAGE
        img_resized = img.resize((width, height), resample)
        
        # ✅ Convert RGBA to RGB if needed for JPEG
        if output_format in ['JPEG', 'JPG']:
            if img_resized.mode == 'RGBA':
                if background:
                    bg = Image.new('RGB', img_resized.size, background)
                    bg.paste(img_resized, mask=img_resized.split()[3] if len(img_resized.split()) > 3 else None)
                    img_resized = bg
                else:
                    # ✅ Convert to RGB with white background
                    bg = Image.new('RGB', img_resized.size, (255, 255, 255))
                    if img_resized.mode == 'RGBA':
                        bg.paste(img_resized, mask=img_resized.split()[3])
                    else:
                        bg.paste(img_resized)
                    img_resized = bg
        
        # ✅ Save resized image to buffer
        buffered = BytesIO()
        save_kwargs = {'format': output_format, 'quality': quality_value, 'optimize': True}
        
        if output_format == 'PNG':
            save_kwargs['compress_level'] = 6
        elif output_format in ['JPEG', 'JPG']:
            save_kwargs['quality'] = quality_value
            save_kwargs['optimize'] = True
            save_kwargs['progressive'] = True
        
        img_resized.save(buffered, **save_kwargs)
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
        
        # ✅ Add watermark for free users
        if should_add_watermark(is_premium):
            watermarked_img = Image.open(BytesIO(buffered.getvalue()))
            watermarked_img = add_watermark_to_image(watermarked_img, is_premium)
            watermarked_buffer = BytesIO()
            watermarked_img.save(watermarked_buffer, format=output_format)
            img_base64 = base64.b64encode(watermarked_buffer.getvalue()).decode()
        
        # ✅ Clean up temp file
        cleanup_temp_files([temp_path])
        
        increment_usage('image_resizer', ip_address)
        current_usage = get_usage_count('image_resizer', ip_address)
        
        return jsonify({
            'success': True,
            'image': f"data:image/{output_format.lower()};base64,{img_base64}",
            'original_size': original_size,
            'new_size': (width, height),
            'usage_count': current_usage,
            'remaining_free': "Unlimited",
            'is_premium': is_premium,
            'max_dimension': max_dimension,
            'format': output_format,
            'quality': quality_value,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"Image Resizer Error: {str(e)}")
        import traceback
        traceback.print_exc()
        cleanup_temp_files([temp_path] if 'temp_path' in locals() else [])
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================
# 12. TEXT TO PDF
# ============================================

@tools_bp.route('/text-to-pdf', methods=['POST', 'OPTIONS'])
def text_to_pdf():
    """Convert Text to PDF (Free: 5 per day, Premium: Unlimited)"""
    if request.method == 'OPTIONS':
        response = jsonify({'success': True})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        ip_address = get_client_ip(request)
        
        # Rate limiting
        if is_rate_limited(ip_address, 'text_to_pdf', limit=20, window=60):
            return jsonify({
                'success': False,
                'error': 'Too many requests. Please try again later.'
            }), 429
        
        track_request(ip_address, 'text_to_pdf')
        
        is_premium = data.get('is_premium', False)
        
        usage_count = get_usage_count('text_to_pdf', ip_address)
        
        if not is_premium and usage_count >= 5:
            return jsonify({
                'success': False,
                'error': 'Free limit reached. Please upgrade to premium for unlimited access.',
                'limit_reached': True,
                'usage_count': usage_count,
                'max_free': 5
            }), 403
        
        text = data.get('text', '')
        title = data.get('title', 'Document')
        
        if not text:
            return jsonify({'success': False, 'error': 'Text is required'}), 400
        
        # Sanitize and limit text
        text = text[:10000]  # Max 10000 chars
        title = sanitize_input(title)[:100]
        
        # Create PDF
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, title[:80])
        
        # Content
        c.setFont("Helvetica", 12)
        y = height - 80
        lines = text.split('\n')
        page_count = 1
        
        for line in lines:
            if y < 50:
                c.showPage()
                page_count += 1
                y = height - 50
                c.setFont("Helvetica", 12)
                # Page header on new page
                c.setFont("Helvetica-Bold", 10)
                c.drawString(50, height - 30, f"{title} - Page {page_count}")
                c.setFont("Helvetica", 12)
            
            # Truncate long lines
            if len(line) > 80:
                line = line[:77] + '...'
            c.drawString(50, y, line)
            y -= 20
        
        # ✅ Add watermark for free users (on last page)
        if should_add_watermark(is_premium):
            c.showPage()
            c.setFont("Helvetica", 10)
            c.setFillColorRGB(0.5, 0.5, 0.5)
            c.drawCentredString(width/2, 50, "Made with ❤️ by Krynova Technologies")
            c.drawCentredString(width/2, 35, "Visit: https://krynovatechnology.pythonanywhere.com")
        
        c.save()
        
        pdf_base64 = base64.b64encode(buffer.getvalue()).decode()
        
        increment_usage('text_to_pdf', ip_address)
        
        return jsonify({
            'success': True,
            'file': pdf_base64,
            'filename': f'{title}.pdf',
            'usage_count': get_usage_count('text_to_pdf', ip_address),
            'remaining_free': max(0, 5 - get_usage_count('text_to_pdf', ip_address)),
            'is_premium': is_premium,
            'has_watermark': should_add_watermark(is_premium)
        })
        
    except Exception as e:
        print(f"Text to PDF Error: {str(e)}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

# ============================================
# ❌ PREMIUM CHECK REMOVED - Now in payment_routes.py
# ❌ PREMIUM SUBSCRIBE REMOVED - Now in payment_routes.py
# ============================================

# ============================================
# HELPER FUNCTIONS
# ============================================

def generate_resume_content(data):
    """Generate professional resume content with support for multiple entries"""
    name = data.get('name', '')
    email = data.get('email', '')
    phone = data.get('phone', '')
    location = data.get('location', '')
    title = data.get('title', '')
    summary = data.get('summary', '')
    skills = data.get('skills', [])
    
    # Handle skills - could be string or array
    if isinstance(skills, str):
        skills = [s.strip() for s in skills.split(',') if s.strip()]
    elif not isinstance(skills, list):
        skills = []
    
    # Get dynamic entries
    experience = data.get('experience', [])
    education = data.get('education', [])
    projects = data.get('projects', [])
    certifications = data.get('certifications', [])
    languages = data.get('languages', [])
    
    # Build resume
    resume = f"""
{name}
{email} | {phone}
{location}

{'-' * 50}

{title}

PROFESSIONAL SUMMARY
{summary if summary else f'Experienced professional with expertise in {", ".join(skills[:3]) if skills else "various technologies"}. Passionate about delivering high-quality solutions.'}

{'-' * 50}

SKILLS
"""
    for skill in skills:
        resume += f"• {skill}\n"
    
    # EXPERIENCE SECTION
    if experience and len(experience) > 0:
        resume += f"""
{'-' * 50}

EXPERIENCE
"""
        for exp in experience:
            if isinstance(exp, dict):
                exp_title = exp.get('title', '')
                company = exp.get('company', '')
                start_date = exp.get('startDate', '')
                end_date = exp.get('endDate', 'Present')
                description = exp.get('description', '')
                
                resume += f"\n{exp_title} | {company}\n"
                resume += f"{start_date} - {end_date}\n"
                if description:
                    # Handle bullet points in description
                    for line in description.split('\n'):
                        line = line.strip()
                        if line:
                            if line.startswith('•') or line.startswith('-'):
                                resume += f"{line}\n"
                            else:
                                resume += f"• {line}\n"
            else:
                # Fallback if it's a string
                resume += f"\n{exp}\n"
    
    # EDUCATION SECTION
    if education and len(education) > 0:
        resume += f"""
{'-' * 50}

EDUCATION
"""
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                start_year = edu.get('startYear', '')
                end_year = edu.get('endYear', 'Present')
                gpa = edu.get('gpa', '')
                description = edu.get('description', '')
                
                resume += f"\n{degree}\n{institution}\n{start_year} - {end_year}"
                if gpa:
                    resume += f"\nGPA: {gpa}"
                if description:
                    resume += f"\n{description}"
                resume += "\n"
            else:
                resume += f"\n{edu}\n"
    
    # PROJECTS SECTION
    if projects and len(projects) > 0:
        resume += f"""
{'-' * 50}

PROJECTS
"""
        for proj in projects:
            if isinstance(proj, dict):
                proj_name = proj.get('name', '')
                year = proj.get('year', '')
                technologies = proj.get('technologies', '')
                description = proj.get('description', '')
                
                resume += f"\n{proj_name}"
                if year:
                    resume += f" ({year})"
                resume += "\n"
                if technologies:
                    resume += f"Technologies: {technologies}\n"
                if description:
                    for line in description.split('\n'):
                        line = line.strip()
                        if line:
                            if line.startswith('•') or line.startswith('-'):
                                resume += f"{line}\n"
                            else:
                                resume += f"• {line}\n"
            else:
                resume += f"\n{proj}\n"
    
    # CERTIFICATIONS SECTION
    if certifications and len(certifications) > 0:
        resume += f"""
{'-' * 50}

CERTIFICATIONS
"""
        for cert in certifications:
            if isinstance(cert, dict):
                cert_name = cert.get('name', '')
                issuer = cert.get('issuer', '')
                year = cert.get('year', '')
                
                resume += f"• {cert_name}"
                if issuer:
                    resume += f" - {issuer}"
                if year:
                    resume += f" ({year})"
                resume += "\n"
            else:
                resume += f"• {cert}\n"
    
    # LANGUAGES SECTION
    if languages and len(languages) > 0:
        resume += f"""
{'-' * 50}

LANGUAGES
"""
        for lang in languages:
            if isinstance(lang, dict):
                language = lang.get('language', '')
                proficiency = lang.get('proficiency', '')
                
                if language:
                    resume += f"• {language}"
                    if proficiency:
                        resume += f" - {proficiency}"
                    resume += "\n"
            else:
                resume += f"• {lang}\n"
    
    return resume

def generate_cover_letter_content(data):
    """Generate professional cover letter"""
    name = data.get('name', '')
    position = data.get('position', '')
    company = data.get('company', '')
    skills = data.get('skills', [])
    experience = data.get('experience', '')
    
    if isinstance(skills, str):
        skills = skills.split(',')
    
    return f"""
{name}
{data.get('email', '')} | {data.get('phone', '')}
{data.get('location', '')}

{datetime.now().strftime('%B %d, %Y')}

Hiring Manager
{company}

Dear Hiring Manager,

I am writing to express my strong interest in the {position} position at {company}. 
With my expertise in {', '.join(skills[:3]) if skills else 'various technologies'}, I am confident I can make significant 
contributions to your team.

Throughout my career, I have developed strong skills in {', '.join(skills) if skills else 'various technologies'}. 
{experience[:200] if experience else 'I am passionate about delivering high-quality work and thrive in collaborative environments.'}

I am particularly drawn to {company} because of your reputation for innovation and 
excellence. I would welcome the opportunity to discuss how my experience aligns 
with the needs of your organization.

Thank you for considering my application.

Sincerely,
{name}
"""

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF"""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
            return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return "Unable to extract text from PDF."

print("✅ tools_routes.py loaded successfully!", flush=True)